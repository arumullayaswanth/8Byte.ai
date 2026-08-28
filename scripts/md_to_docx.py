import sys
from docx import Document
from docx.shared import Pt

src = sys.argv[1]
dst = sys.argv[2]

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

with open(src, encoding="utf-8") as f:
    lines = f.read().splitlines()

para_buf = []

def flush_para():
    if para_buf:
        doc.add_paragraph(" ".join(para_buf))
        para_buf.clear()

for line in lines:
    stripped = line.strip()
    if stripped.startswith("# "):
        flush_para()
        doc.add_heading(stripped[2:], level=0)
    elif stripped.startswith("## "):
        flush_para()
        doc.add_heading(stripped[3:], level=1)
    elif stripped.startswith("### "):
        flush_para()
        doc.add_heading(stripped[4:], level=2)
    elif stripped.startswith("- "):
        flush_para()
        doc.add_paragraph(stripped[2:], style="List Bullet")
    elif stripped == "":
        flush_para()
    else:
        para_buf.append(stripped)

flush_para()
doc.save(dst)
print("wrote", dst)
